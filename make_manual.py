# -*- coding: utf-8 -*-
"""產生「譯彩紛呈：重譯文本分析系統」使用說明手冊 .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK = "Microsoft JhengHei"   # 微軟正黑體
NAVY = RGBColor(0x0E, 0x15, 0x30)
BLUE = RGBColor(0x25, 0x63, 0xEB)
PURPLE = RGBColor(0x7C, 0x3A, 0xED)
GREEN = RGBColor(0x0F, 0x9D, 0x58)
GREY = RGBColor(0x47, 0x55, 0x69)

doc = Document()

# ── 全文預設字型（含中文 eastAsia）──
def _apply_cjk(style):
    style.font.name = CJK
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), CJK)
    rfonts.set(qn("w:ascii"), CJK)
    rfonts.set(qn("w:hAnsi"), CJK)

normal = doc.styles["Normal"]
normal.font.size = Pt(11)
_apply_cjk(normal)
for sname in ("Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle",
              "List Bullet", "List Number"):
    try:
        _apply_cjk(doc.styles[sname])
    except KeyError:
        pass


def set_run(r, size=11, color=None, bold=False):
    r.font.name = CJK
    r._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
    r.font.size = Pt(size)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color


def para(text="", size=11, color=None, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_run(p.add_run(text), size, color, bold)
    return p


def h1(text):
    p = doc.add_heading(level=1)
    set_run(p.add_run(text), 17, NAVY, True)
    return p


def h2(text):
    p = doc.add_heading(level=2)
    set_run(p.add_run(text), 14, BLUE, True)
    return p


def bullet(text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    set_run(p.add_run(text), size)
    return p


def numbered(text, size=11):
    p = doc.add_paragraph(style="List Number")
    set_run(p.add_run(text), size)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].clear()
        set_run(c.paragraphs[0].add_run(htext), 10.5, RGBColor(0xFF, 0xFF, 0xFF), True)
        # 表頭底色
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "2563EB")
        c._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].clear()
            set_run(cells[i].paragraphs[0].add_run(str(val)), 10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ══════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()
para("譯彩紛呈：重譯文本分析系統", 28, NAVY, True, WD_ALIGN_PARAGRAPH.CENTER, 4)
para("TransPrism — Retranslation I　語料淬煉 Corpus Refinery",
     14, PURPLE, True, WD_ALIGN_PARAGRAPH.CENTER, 24)
para("系統使用說明手冊", 20, BLUE, True, WD_ALIGN_PARAGRAPH.CENTER, 30)
para("以中央研究院 CKIP Transformers 為核心的中文斷詞、詞性標註，"
     "並結合 NER + LLM 進行專名探勘與字典擴充。",
     11.5, GREY, False, WD_ALIGN_PARAGRAPH.CENTER, 8)
para("版本：1.0　　文件產生日期：2026-06-25",
     10.5, GREY, False, WD_ALIGN_PARAGRAPH.CENTER, 4)
doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. 系統簡介
# ══════════════════════════════════════════════════════════
h1("一、系統簡介")
para("本系統是一套在本機端執行的中文文本分析工具，透過瀏覽器操作（Gradio 介面），"
     "主要提供兩大功能：")
bullet("斷詞與詞性標註：將中文原文切分為詞彙，並標註每個詞的詞性（POS），"
       "支援自訂字典與多項斷詞前後處理規則。")
bullet("專名探勘（擴充字典）：以命名實體辨識（NER）找出文本中的專有名詞候選，"
       "再交由大型語言模型（LLM）判斷是否為真正專名，協助擴充斷詞字典。")
para("系統具備 GPU 自動加速（偵測到 NVIDIA 顯示卡時自動啟用），"
     "並以全域互斥鎖確保兩項作業不會同時佔用運算資源。", space_after=10)

# ══════════════════════════════════════════════════════════
# 2. 系統需求
# ══════════════════════════════════════════════════════════
h1("二、系統需求與執行環境")
table(
    ["項目", "需求 / 建議"],
    [
        ["作業系統", "Windows 10 / 11（亦支援 macOS、Linux）"],
        ["Python", "3.9 以上（本機實測 3.12）"],
        ["記憶體", "建議 8GB 以上"],
        ["顯示卡（選用）", "NVIDIA GPU（本機為 RTX 4090，24GB）可大幅加速"],
        ["核心套件", "ckip-transformers、gradio、torch（CUDA 12.4）、transformers"],
        ["LLM（專名探勘用）", "需任一供應商 API Key：OpenRouter／OpenAI／Gemini／DeepSeek／Anthropic／Groq"],
    ],
    widths=[1.8, 4.6],
)

# ══════════════════════════════════════════════════════════
# 3. 安裝與啟動
# ══════════════════════════════════════════════════════════
h1("三、安裝與啟動")
h2("3.1 安裝步驟")
numbered("取得程式：git clone 專案，或下載原始碼到本機資料夾。")
numbered("建立虛擬環境：python -m venv venv，並啟用之。")
numbered("安裝 PyTorch（GPU 版）：pip install torch torchvision torchaudio "
         "--index-url https://download.pytorch.org/whl/cu124")
numbered("安裝其餘套件：pip install -r requirements.txt")
numbered("（專名探勘需要）設定 .env 檔，填入 API Key 與模型名稱。")
h2("3.2 啟動系統")
numbered("執行：python app.py")
numbered("待主控台顯示啟動完成後，於瀏覽器開啟：http://127.0.0.1:7860")
para("提示：首次啟動會自動下載 CKIP 模型（約需數分鐘），請保持網路連線。",
     10.5, GREY, space_after=10)

# ══════════════════════════════════════════════════════════
# 4. 系統架構圖
# ══════════════════════════════════════════════════════════
h1("四、系統架構")
para("下圖呈現系統的整體分層與兩大流程的處理步驟，"
     "以及底層的模型與運算資源（CKIP 模型、GPU、LLM API）：")
doc.add_picture(r"D:\ckip_segmentation_sumin\系統架構圖.png", width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para("（圖）譯彩紛呈：重譯文本分析系統 — 系統架構圖",
     9.5, GREY, False, WD_ALIGN_PARAGRAPH.CENTER, 12)

# ══════════════════════════════════════════════════════════
# 5. 功能說明 A：斷詞
# ══════════════════════════════════════════════════════════
h1("五、功能說明（一）：斷詞與詞性標註")
h2("5.1 操作步驟")
numbered("切換至「斷詞與詞性標註」分頁。")
numbered("於「上傳文本檔案」上傳一個或多個 .txt 原始文本（UTF-8 編碼）。")
numbered("（選用）於「上傳自訂字典」上傳字典檔（.txt，每行一個詞彙）。")
numbered("點擊「開始斷詞與標註」，右側「處理紀錄」會即時顯示進度。")
numbered("完成後於「下載結果（ZIP）」下載打包好的結果檔。")
h2("5.2 處理流程與修正規則")
bullet("斷詞前處理：半形標點轉全形；異體字／錯字修正（如「内→內」「爲→為」）。")
bullet("斷詞與標註：以 CKIP（bert-base）進行 WS 斷詞與 POS 詞性標註。")
bullet("自訂字典合併：以最長匹配將連續詞彙合併為字典詞。")
bullet("斷詞後修正：套用內建規則（i–liv）修正特定人名／地名／詞性，"
       "並刪除多餘空白、段落開頭多餘符號等。")
bullet("整行刪除：完成斷詞後，刪除「以 _ETCCATEGORY 詞性開頭」的整行"
       "（例如以省略號開頭的行）。")
h2("5.3 輸出格式")
para("輸出為「詞彙_詞性」以空白分隔，例如：那_Nep 一陣子_Nd 東京都_Nc；"
     "每個輸入檔對應一個 _seg.txt，最後統一打包為 ZIP。")
h2("5.4 防呆機制")
para("若上傳的檔案其實已是「斷詞輸出」格式（詞_詞性），系統會立即停止並提示"
     "「請改上傳原始、未斷詞的文本」，避免重複斷詞造成亂碼。", space_after=10)

# ══════════════════════════════════════════════════════════
# 6. 功能說明 B：專名探勘
# ══════════════════════════════════════════════════════════
h1("六、功能說明（二）：專名探勘（擴充字典）")
h2("6.1 操作步驟")
numbered("切換至「專名探勘（擴充字典）」分頁。")
numbered("上傳文本檔案；（選用）上傳現有字典（用於排除已收錄詞並作為匯出基底）。")
numbered("選擇「LLM 供應商」，填入「API Key」（留空則使用 .env 設定）。")
numbered("選擇「模型」（可從清單挑選，或選『自行輸入模型名稱…』後直接輸入）。")
numbered("點擊「開始專名探勘」，右上方「進度」條與「處理紀錄」會即時回報。")
numbered("於候選表格勾選要「收錄」的詞，點「匯出選取詞典」，再下載更新後字典。")
h2("6.2 支援的 LLM 供應商")
table(
    ["供應商", "說明 / 常用模型"],
    [
        ["OpenRouter", "聚合多家模型（gemma、gemini、gpt、claude、deepseek…）"],
        ["OpenAI", "gpt-4o-mini、gpt-4o、gpt-4.1…"],
        ["Google Gemini", "gemini-2.0-flash、gemini-1.5-pro…"],
        ["DeepSeek", "deepseek-chat、deepseek-reasoner"],
        ["Anthropic", "claude-3-5-haiku／sonnet、claude-sonnet-4"],
        ["Groq", "llama-3.3-70b、llama-3.1-8b"],
    ],
    widths=[1.8, 4.6],
)
para("進度與即時回報：因 LLM 回應速度時快時慢，系統於每批審核前後更新進度條，"
     "並顯示已用時間與已審查／建議收錄數量，方便確認系統仍在運作。", space_after=10)

# ══════════════════════════════════════════════════════════
# 7. .env 設定
# ══════════════════════════════════════════════════════════
h1("七、.env 設定說明")
para("專名探勘所需的 API Key 與模型，可寫在專案根目錄的 .env 檔（介面欄位留空時自動採用）：")
table(
    ["設定項", "說明", "範例"],
    [
        ["OPENROUTER_API_KEY", "API 金鑰（亦相容 LLM_API_KEY）", "sk-or-v1-xxxx"],
        ["OPENROUTER_MODEL", "預設模型（亦相容 LLM_MODEL）", "google/gemma-4-26b-a4b-it"],
    ],
    widths=[2.2, 2.6, 1.8],
)
para("安全提醒：.env 已被 .gitignore 排除，不會上傳到 GitHub，API Key 不會外洩。"
     "請勿將 Key 直接寫進程式碼或提交至版本庫。", 10.5, RGBColor(0xB0, 0x20, 0x20), True,
     space_after=10)

# ══════════════════════════════════════════════════════════
# 8. 注意事項 / FAQ
# ══════════════════════════════════════════════════════════
h1("八、注意事項與常見問題")
table(
    ["問題 / 情境", "說明與處理方式"],
    [
        ["兩項作業可同時跑嗎？", "不行。系統以互斥鎖限制一次只執行一項，"
                                "避免顯存不足或互相拖慢；忙碌時會提示稍候。"],
        ["上傳已斷詞檔做斷詞？", "斷詞分頁會停止並要求重新上傳原文；"
                                "專名探勘分頁則會自動還原為原始文字後繼續。"],
        ["速度很慢？", "確認已啟用 GPU（介面上方會顯示「運算裝置：GPU…」）；"
                       "專名探勘速度另受 LLM 供應商回應速度影響。"],
        ["找不到 API key？", "於介面 API Key 欄位填入，或設定 .env 的 LLM_API_KEY。"],
        ["首次啟動很久？", "首次需下載 CKIP 模型，屬正常現象，完成後即快速。"],
    ],
    widths=[1.9, 4.5],
)

# ══════════════════════════════════════════════════════════
# 9. 版本與授權
# ══════════════════════════════════════════════════════════
h1("九、版本與授權")
bullet("核心斷詞模型：中央研究院 CKIP Transformers（GPL-3.0 授權）。")
bullet("介面框架：Gradio；深度學習框架：PyTorch（CUDA 12.4）。")
bullet("本手冊對應版本：1.0。")

out = r"D:\ckip_segmentation_sumin\系統使用說明手冊.docx"
doc.save(out)
print("已輸出:", out)
